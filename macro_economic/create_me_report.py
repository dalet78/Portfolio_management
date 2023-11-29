import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.dates import DateFormatter
import matplotlib.dates as mdates
import os
import datetime
# import textwrap
from docx import Document
from docx.shared import Inches, Pt


class CreateReport():
    def __init__(self, config_dict, country):
        self.report_dict = config_dict
        self.country_data = self.report_dict['countries'][country]
        self.me_report_path = "macro_economic/report"
        self.date = datetime.datetime.now().strftime("%Y-%m-%d")
        self.country = country
        self.pdf_file = ""
        self.data_path = "macro_economic/Data"
        
        
    def _pdf_name_report(self):    
        if self.coutry is not None:
            self.pdf_file = f"{self.me_report_path}/{self.date}_{self.country}_Report.pdf"
        else:  
            self.pdf_file = f"{self.me_report_path}/{self.date}_all_countries_Report.pdf"
    
    def _create_graph(self, csv_file, description):
        # Leggi i dati dal file CSV
        data = pd.read_csv(csv_file)

        # Assumi che la prima colonna contenga le date o etichette per l'asse X
        # e che sia nel formato 'YYYY-MM-DD'. Modifica se necessario.
        data['Date'] = pd.to_datetime(data.iloc[:, 0])  
        data.set_index('Date', inplace=True)

        # Crea un grafico in base ai dati (modifica questa parte in base al tipo di grafico desiderato)
        plt.figure()
        data.plot()  # Esempio di plot generico
        plt.title(csv_file)

        # Formatta l'asse X per visualizzare le date
        plt.gca().xaxis.set_major_formatter(DateFormatter('%Y-%m-%d'))
        plt.gca().xaxis.set_major_locator(mdates.YearLocator())
        plt.xticks(rotation=45)

        
        # Ottieni gli ultimi 5 valori con date
        last_five = data.iloc[-5:]
        text_str = last_five.to_string(header=False)

        # Aggiungi questi valori come testo nel grafico
        plt.gcf().text(0.5, -0.1, text_str, ha='center')

        directory = os.path.dirname(csv_file)
        if not os.path.exists(directory):
            os.makedirs(directory)

        # Salva il grafico come immagine
        image_file = csv_file.replace('.csv', '.png')
        plt.savefig(image_file, bbox_inches='tight')
        plt.close()
        if not os.path.exists(image_file):
            print(f"Error: Image file {image_file} does not exist.")
        return image_file

    def report_create(self):
        document = Document()

        # Aggiungi un titolo al documento
        title = f"Rapporto Macroeconomico {self.country} del {self.date} "  # Modifica con il titolo desiderato
        title_paragraph = document.add_paragraph(title, 'Title')
        # Imposta la dimensione del carattere del titolo
        for run in title_paragraph.runs:
            run.font.size = Pt(14)

        # graphs_info = {f"{os.path.join(self.data_path, self.country)}/{key}.csv": value for key, value in self.report_dict["indicator_description"].items()}
        # table = document.add_table(rows=1, cols=2)
        # current_row = table.rows[0].cells
        # cell_index = 0

        # for index, (file_name, description) in enumerate(graphs_info.items()):
        #     image_file = self._create_graph(file_name, description)

        #     # Se hai raggiunto la seconda colonna, aggiungi una nuova riga
        #     if cell_index == 2:
        #         current_row = table.add_row().cells
        #         cell_index = 0

        #     # Aggiungi il grafico e la descrizione alla cella corrente
        #     paragraph = current_row[cell_index].paragraphs[0]
        #     run = paragraph.add_run()
        #     run.add_picture(image_file, width=Inches(3))  # Imposta la larghezza in base alla dimensione desiderata

        #     # Aggiungi una nuova linea e poi la descrizione
        #     run = paragraph.add_run("\n")
        #     run = paragraph.add_run(description)
        #     font = run.font
        #     font.size = Pt(9)


        #     # Passa alla cella successiva
        #     cell_index += 1

        #     # Rimuovi il file immagine
        #     os.remove(image_file)

        # # Salva il documento Word
        # word_file_name = 'final_output.docx'
        # document.save(word_file_name)

        # print(f"Documento Word creato con successo: '{word_file_name}'")

        for category, indicators in self.country_data['indicators'].items():
            # Aggiungi un'intestazione per ogni categoria
            document.add_paragraph(category, 'Heading1')

            # Elabora gli indicatori
            if isinstance(indicators, dict):                
                for indicator_name, indicator_code in indicators.items():
                    self._add_indicator_to_document(document, indicator_name, indicator_code)
            else:
                self._add_indicator_to_document(document, category, indicators)
        
        # Salva il documento Word
        word_file_name = f'{self.me_report_path}/{self.date}_{self.country}_Report.docx'
        document.save(word_file_name)
        print(f"Documento Word creato con successo: '{word_file_name}'")

    def _add_indicator_to_document(self, document, indicator_name, indicator_code):
        # Implementa la logica per creare il grafico e aggiungere la descrizione
        csv_file = os.path.join(self.data_path, self.country, f"{indicator_code}.csv")
        description = f"Descrizione per {indicator_name}"
        image_file = self._create_graph(csv_file, description)

        # Aggiungi il grafico e la descrizione al documento
        document.add_paragraph(indicator_name, 'Heading2')
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_file, width=Inches(3))
        paragraph.add_run(description)

        # Rimuovi il file immagine se necessario
        os.remove(image_file)
